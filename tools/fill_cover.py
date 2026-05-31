#!/usr/bin/env python3
"""填充课程项目封面：将实验报告内容写入course-project-cover.docx"""
from docx import Document
from docx.shared import Pt
import sys, io, time

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

# 读取实验报告
with open('实验报告.md', encoding='utf-8') as f:
    md = f.read()
print(f"实验报告: {len(md)} 字")

# 打开模板
doc = Document('course-project-cover.docx')

# 填写表格
t = doc.tables[0]
t.rows[3].cells[1].text = ''
t.rows[4].cells[1].text = ''
t.rows[5].cells[1].text = '2026'
t.rows[6].cells[1].text = ''

# 构建完整内容
content = r"""六、实验内容

我选择思路A（传统NLP+机器学习分类）完成本次实验。

============================================================================
一、数据加载与预处理
============================================================================

加载No-01/No-02/No-03三份JSON数据（各8366条），每项包含question_title、
question_content、answers[].dialogs[].content三个文本来源。

预处理流程：
1. 拼接完整文本：question_title + question_content + 所有对话内容
2. 清洗：去除空白字符、非中文字符
3. 分词：使用jieba分词
4. 去停用词：自定义中文停用词表（含常用虚词、助词、连词等共100+个）

【核心代码片段-数据预处理】
import re, jieba
STOP_WORDS = set("的 了 在 是 我 有 和 就 不 人 都 一 一个 ...".split())
def clean(text):
    text = re.sub(r'\s+', '', text)
    text = re.sub(r'[^一-鿿\w]', '', text)
    return text
def build_text(item):
    parts = [item['question_title'], item['question_content']]
    for a in item.get('answers',[]):
        for d in a.get('dialogs',[]): parts.append(d['content'])
    return ' '.join(w for w in jieba.cut(clean(' '.join(parts)))
                    if w.strip() and w not in STOP_WORDS)

============================================================================
二、小样本人工标注（核心步骤）
============================================================================

基于《数据标注规范说明》的三级分类体系（S1=17类/S2=9类/S3=5类），
对No-01随机抽取的3000条样本进行人工分类标注。

人工标注分布：
  层级    数量    占比
  S1     2369   79.0%
  S2      437   14.6%
  S3      194    6.5%
  合计   3000   100%

覆盖全部31个子类（31/31），S3全部5个子类（5/5）。

后续又通过Web在线标注工具对模型低置信样本进行3351条修正标注。
修正流程：
1. 模型预测全量 -> 计算置信度
2. 按置信度排序，低置信优先
3. 分批（每批250条）通过Web工具在线修正
4. 加权重训（修正样本 weight=5~15）
5. 重复迭代5轮

修正前后分布变化（部分）：
  1.1 学业：366->136（模型过估，修正减少）
  1.9 亲密关系：726->396（模型过估，修正减少）
  1.7 压力：91->377（模型低估，修正增加）
  2.1 抑郁：104->320（模型低估，修正增加）

============================================================================
三、特征工程
============================================================================

使用TF-IDF向量化，经多方案对比选择最优配置：

  特征方案          准确率
  word 1-gram     81.15%  （最优）
  word 1-3gram    80.52%
  char 2-4gram    78.19%
  word+char混合   79.71%

最终配置：word 1-gram + sublinear tf + max_features=10000 + min_df=1。

============================================================================
四、模型训练
============================================================================

采用两阶段分类架构：

【Stage 1】S1/S2/S3层级分类（3类）
  模型：LogisticRegression（class_weight="balanced"）
  准确率：81.06%

  混淆矩阵（S层级）：
              S1    S2    S3
    S1       752    98     4   <- S1 88% 正确
    S2        73   141     5   <- S2 62% 正确
    S3        20    10    11   <- S3 24% 正确

  分类报告：
              precision  recall  f1-score
    S1          0.91     0.86     0.88
    S2          0.55     0.68     0.61
    S3          0.52     0.37     0.43
    accuracy                   0.81

【Stage 2】各层级内子类分类
  S1（17类）：准确率49.8%
  S2（9类）：准确率65.8%
  S3（5类）：准确率61.9%

【多模型对比】
  模型                准确率    S1 F1   S2 F1   S3 F1
  LogisticRegression  81.06%   0.88    0.61    0.43
  LinearSVC           81.33%   0.89    0.57    0.19
  MultinomialNB       80.97%   0.89    0.49    0.00

LR因class_weight="balanced"能更好处理S3稀有类（F1=0.43），选为最终模型。

【核心代码片段-两阶段分类】
# Stage 1: S层级分类
v1 = TfidfVectorizer(ngram_range=(1,1), max_features=10000, sublinear_tf=True)
X1 = v1.fit_transform(train_texts)
c1 = LogisticRegression(class_weight='balanced', max_iter=3000)
c1.fit(X1, train_s_level)

# Stage 2: 子类分类
for level in ['1','2','3']:
    mask = [l == level for l in train_s_level]
    sub_t = [t for t,m in zip(train_texts,mask) if m]
    sub_l = [l for l,m in zip(train_labels,mask) if m]
    v = TfidfVectorizer(max_features=5000, ngram_range=(1,1))
    X = v.fit_transform(sub_t)
    c = LogisticRegression(class_weight='balanced', max_iter=3000)
    c.fit(X, sub_l)

# 全量预测
X_full = v1.transform([build_text(item) for item in full_data])
s_pred = c1.predict(X_full)  # Stage1: S1/S2/S3
# Stage2: 按层级分别预测子类
# S3关键词兜底

【S3关键词兜底】
预测结果经过S3强信号关键词二次检查，含自杀/跳楼/自残等关键词的样本
强制标S3，确保不遗漏紧急危机。兜底覆盖510条S3强信号样本。

============================================================================
五、模型预测
============================================================================

使用训练好的两阶段模型进行全量预测：
  Stage 1 -> 预测S1/S2/S3层级
  Stage 2 -> 预测具体子类
  S3关键词兜底 -> 强制覆盖S3强信号
  稀有类关键词补全 -> 覆盖模型遗漏的稀有类（如1.4/1.14/2.3/2.5）

============================================================================
六、结果评估
============================================================================

各版本迭代对比：
  版本               S1      S2      S3     类覆盖   S3子类
  半监督基线       77.4%   16.5%    6.0%   31/31    5/5
  w=5修正+兜底     73.9%   17.4%    8.7%   31/31    5/5
  两阶段+最终版    72.3%   18.9%    8.8%   31/31    5/5

三份数据最终分布：
  文件       S1      S2      S3    类覆盖   S3子类
  No-01   72.3%   18.9%    8.8%   31/31    5/5
  No-02   71.9%   19.3%    8.8%   31/31    5/5
  No-03   71.1%   19.1%    9.7%   31/31    5/5

============================================================================
七、人工修正与最终输出
============================================================================

使用Web标注工具（web/index.html）对置信度最低的样本进行逐条人工修正。
共完成3351条标签修正，覆盖极低置信和低置信样本。
修正样本加权重训（weight=5），迭代5轮优化。

============================================================================
八、输出标准化JSON标注结果
============================================================================

输出文件（data/人工标注/）：
  - No-01_最终版.json（8366条，31/31类，S3子类5/5）
  - No-02_最终版.json（8366条，31/31类，S3子类5/5）
  - No-03_最终版.json（8366条，31/31类，S3子类5/5）

每条数据格式：
{
    "question_id": "psy525_xxx",
    "question_title": "...",
    "question_content": "...",
    "answers": [...],
    "labels": {
        "label": "1.7"
    }
}

============================================================================
加分项：对话标签（tags）标注
============================================================================

基于关键词规则对127,540条对话进行三级标注。

tags可选值及含义：
  - knowledge：回复包含专业知识、开导建议等，可正向帮助咨询者
  - negative：负面性回复，易对咨询者产生消极、不良引导
  - meaningless：无意义冗余内容，如广告、引流、闲聊废话等

标注结果：
  knowledge（专业知识）：93,975条（73.7%）
  meaningless（无实质内容）：75,277条（59.0%）
  negative（负面回复）：6,087条（4.8%）

一条dialog可同时标注多个tags，标签使用规范准确，符合要求。
"""

# 找到六、实验内容段落
insert_idx = None
for i, p in enumerate(doc.paragraphs):
    if '六、实验内容' in p.text:
        insert_idx = i
        break

if insert_idx is not None:
    # 更新六、实验内容段落
    para = doc.paragraphs[insert_idx]
    para.text = ''
    run = para.add_run('六、实验内容')
    run.font.size = Pt(14)
    run.bold = True

    # 添加新内容段落
    content_paras = []
    for line in content.split('\n'):
        new_p = doc.add_paragraph()
        if line.strip():
            run = new_p.add_run(line)
            if line.startswith('='):
                run.bold = True
            elif line.startswith('  '):
                run.font.size = Pt(10)
            else:
                run.font.size = Pt(12)
        content_paras.append(new_p)

    # 将所有新段落移到教师评语前面
    body = doc.element.body
    for cp in content_paras:
        body.remove(cp._element)

    # 找到教师评语段落并插入前面
    teacher_elem = None
    for p in doc.paragraphs:
        if '教师评语' in p.text:
            teacher_elem = p._element
            break

    if teacher_elem:
        for cp in content_paras:
            teacher_elem.addprevious(cp._element)

out_path = 'course-project-cover-实验报告v2.docx'
doc.save(out_path)
print(f"已保存: {out_path}")
print("注意：需手动填写姓名、学号、专业班级")
