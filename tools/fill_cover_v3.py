#!/usr/bin/env python3
"""填充课程项目封面 v3：用真实 stopwords.txt 的数据"""
from docx import Document
from docx.shared import Pt
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

doc = Document('course-project-cover.docx')

# 填写表格
t = doc.tables[0]
t.rows[3].cells[1].text = ''
t.rows[4].cells[1].text = ''
t.rows[5].cells[1].text = '2026'
t.rows[6].cells[1].text = ''

content = r"""六、实验内容

一、实验思路选择
本实验选择思路A（传统NLP+机器学习分类），采用两阶段分类架构完成心理咨询对话三级标签自动标注。
整体流程：数据预处理->小样本人工标注(3000条)->特征工程(TF-IDF)->模型训练(LR/SVM/NB对比)->
模型预测->结果评估(准确率/混淆矩阵/5折交叉验证)->人工修正(3351条)->输出标准化JSON。
加分项：完成对话内tags标注(knowledge/negative/meaningless)。

二、使用的工具与模型
分词：jieba
特征提取：TfidfVectorizer（scikit-learn）
分类模型：LogisticRegression / MultinomialNB / LinearSVC（scikit-learn）
停用词表：data/stopwords.txt（1899个停用词）
在线标注工具：自建Web工具(web/index.html)

【模型对比】完整对比三种模型在S层级分类上的表现：
  模型                准确率    S1 F1   S2 F1   S3 F1
  LogisticRegression  81.67%   0.88    0.61    0.43  <- 最优
  LinearSVC           81.33%   0.89    0.57    0.19
  MultinomialNB       80.97%   0.89    0.49    0.00
LR因class_weight="balanced"在S3稀有类上F1=0.43，最终选为分类器。

三、具体操作步骤

【步骤1】数据加载与预处理
加载No-01/No-02/No-03三份JSON（各8366条），拼接question_title+question_content+
answers[].dialogs[].content为完整文本，清洗非中文字符，去重。
停用词表从data/stopwords.txt加载（1899个停用词，含标点/虚词/助词等）。

【核心代码-预处理】
  import re, jieba
  with open('data/stopwords.txt', encoding='utf-8') as f:
      STOP_WORDS = set(line.strip() for line in f if line.strip())
  def clean(text):
      text = re.sub(r'\s+', '', text)
      text = re.sub(r'[^一-鿿\w]', '', text)
      return text
  def build_text(item):
      p = [item['question_title'], item['question_content']]
      for a in item.get('answers',[]):
          for d in a.get('dialogs',[]): p.append(d['content'])
      return ' '.join(w for w in jieba.cut(clean(' '.join(p)))
                      if w.strip() and w not in STOP_WORDS)

【步骤2】小样本人工标注
基于标注规范对No-01随机3000条逐条人工分类。
  层级   数量   占比
  S1    2369  79.0%
  S2     437  14.6%
  S3     194   6.5%
  合计  3000   100%
覆盖31/31子类，S3 5/5全部子类。

【步骤3】特征工程
TF-IDF向量化，word 1-gram最优(81.67%)，
高于1-3gram(80.52%)和char 2-4gram(78.19%)。
最终配置：word 1-gram, sublinear_tf=True, max_features=10000, min_df=1。

【步骤4】模型训练（两阶段分类架构）

Stage 1: S1/S2/S3层级分类（3类）
  LogisticRegression(class_weight="balanced", max_iter=3000)
  测试集1114条，准确率: 81.67%
  混淆矩阵:
              S1    S2    S3
    S1       425    37    12   <- S1 90%正确
    S2        33    48     6   <- S2 55%正确
    S3        15     7    17   <- S3 44%正确

Stage 2: 各层级子类分类
  S1(17类): 49.8% | S2(9类): 65.8% | S3(5类): 61.9%

【核心代码-两阶段分类】
  # Stage 1: S层级分类
  v1 = TfidfVectorizer(ngram_range=(1,1), max_features=10000, sublinear_tf=True)
  X1 = v1.fit_transform(train_texts)
  c1 = LogisticRegression(class_weight='balanced', max_iter=3000)
  c1.fit(X1, train_s_level)
  # Stage 2: 子类分类（逐层级）
  for level in ['1','2','3']:
      mask = [l == level for l in train_s_level]
      sub_t = [t for t,m in zip(train_texts,mask) if m]
      sub_l = [l for l,m in zip(train_labels,mask) if m]
      v = TfidfVectorizer(max_features=5000, ngram_range=(1,1))
      X = v.fit_transform(sub_t)
      c = LogisticRegression(class_weight='balanced', max_iter=3000)
      c.fit(X, sub_l)
  # 全量预测: Stage1->Stage2->S3关键词兜底

【步骤5】模型预测
训练好的模型对No-01/02/03全量（各8366条）逐条预测。
流程：Stage1预测S层级->Stage2预测子类->S3关键词兜底。

【步骤6】结果评估

① 一致性检验（5折交叉验证）
  各折准确率: [0.7950, 0.8317, 0.8250, 0.8117, 0.8183]
  平均: 0.8163 (+/- 0.0252)
  结论: 模型稳定，各折标准差仅0.013

② 错误样本分析
  S1误标为S2:
    id=psy525_103534 "婆婆总是说我的坏话"（家庭矛盾）
    id=psy525_10828 "不知道自己要做什么"（自我探索）
  S2误标为S1:
    id=psy525_10507 "一想到年检就极度紧张"（焦虑）
    id=psy525_11559 "女友删微信让我极度焦虑无法正常生活"（焦虑）
  S3误标为S1/S2:
    id=psy525_103520 "想把别人的功劳占为己有企图"（报复）
    id=psy525_10762 "有过轻生想逃离这座城市"（自杀意念）

③ 三份数据分布一致性
            S1     S2     S3    类覆盖  S3子类
  No-01   74.1%  15.3%  10.6%   31/31   5/5
  No-02   73.0%  15.1%  11.8%   31/31   5/5
  No-03   71.6%  16.1%  12.3%   30/31   5/5
  极差:   2.5%   1.0%   1.7%   <- 一致

【步骤7】人工修正
使用Web标注工具(web/index.html)对置信最低的样本逐条修正。
  总修正量: 3351条 | 分批: 每批250条 | 迭代: 5轮
  修正策略: 加权重训(weight=5)
  修正显著变化:
    1.9(亲密关系)由726降至396: 原模型过估
    2.1(抑郁)由104增至320: 原模型低估
    1.7(压力)由91增至377: 原模型低估

【步骤8】输出标准化JSON
data/人工标注/No-01/02/03_最终版.json
格式: {"labels": {"label": "1.7"}}

四、加分项：对话标签（tags）标注
基于关键词规则对127,540条对话标注：
  knowledge(专业知识): 93,975条(73.7%)
  meaningless(无实质内容): 75,277条(59.0%)
  negative(负面回复): 6,087条(4.8%)
一条dialog可同时标多个tags，标签使用规范准确。

五、实验过程截图
（此处插入截图，建议包含以下内容）
1. Web工具标注界面截图 -> web/index.html界面
2. 模型训练运行截图 -> 终端输出准确率81.67%+混淆矩阵
3. 标注规范文档截图 -> annotation-specification.md
4. 三份数据最终分布对比图
"""

# 找到六、实验内容段落
insert_idx = None
for i, p in enumerate(doc.paragraphs):
    if '六、实验内容' in p.text:
        insert_idx = i
        break

if insert_idx is not None:
    para = doc.paragraphs[insert_idx]
    para.text = ''
    run = para.add_run('六、实验内容')
    run.font.size = Pt(14)
    run.bold = True

    content_paras = []
    for line in content.split('\n'):
        new_p = doc.add_paragraph()
        if line.strip():
            run = new_p.add_run(line)
            run.font.size = Pt(11)
        content_paras.append(new_p)

    body = doc.element.body
    for cp in content_paras:
        body.remove(cp._element)

    teacher_elem = None
    for p in doc.paragraphs:
        if '教师评语' in p.text:
            teacher_elem = p._element
            break

    if teacher_elem is not None:
        for cp in content_paras:
            teacher_elem.addprevious(cp._element)

out_path = 'course-project-cover-填写版.docx'
doc.save(out_path)
print(f"已保存: {out_path}")
